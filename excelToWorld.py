import os
import shutil
import glob
from openpyxl import load_workbook
from docx import Document
from pathlib import Path
import re

class DocumentProcessor:
    """文档处理类：自动处理所有文档和Excel文件"""
    
    def __init__(self):
        self.output_dir = "生成文档"
        
    def find_docx_files(self):
        """查找当前目录下的所有.docx文档"""
        docx_files = glob.glob("*.docx")
        # 排除以~$开头的临时文件
        docx_files = [f for f in docx_files if not f.startswith('~$')]
        return docx_files
    
    def find_excel_files(self):
        """查找当前目录下的所有.xlsx文件"""
        excel_files = glob.glob("*.xlsx")
        # 排除以~$开头的临时文件
        excel_files = [f for f in excel_files if not f.startswith('~$')]
        return excel_files
    
    def validate_files(self):
        """验证必要的文件是否存在"""
        docx_files = self.find_docx_files()
        excel_files = self.find_excel_files()
        
        if not docx_files:
            raise FileNotFoundError("当前目录下未找到.docx文档")
        
        if not excel_files:
            raise FileNotFoundError("当前目录下未找到.xlsx文件")
        
        print(f"✅ 找到 {len(docx_files)} 个Word文档: {docx_files}")
        print(f"✅ 找到 {len(excel_files)} 个Excel文件: {excel_files}")
        
        return docx_files, excel_files
    
    def create_output_directory(self):
        """创建输出目录"""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            print(f"📁 创建输出目录: {self.output_dir}")
    
    def read_excel_data(self, excel_file):
        """从Excel文件读取所有数据"""
        try:
            workbook = load_workbook(excel_file)
            sheet = workbook.active
            
            # 读取表头（第一行）
            headers = {}
            for col in range(1, sheet.max_column + 1):
                header_cell = sheet.cell(row=1, column=col)
                if header_cell.value:
                    headers[col] = str(header_cell.value).strip()
            
            if not headers:
                raise ValueError(f"Excel文件 '{excel_file}' 中未找到表头")
            
            print(f"📊 从 '{excel_file}' 读取到表头: {list(headers.values())}")
            
            # 读取所有数据行（从第二行开始到末尾）
            names = []
            data_rows = []
            total_rows = 0
            
            for row in range(2, sheet.max_row + 1):
                # 检查第一列是否有数据
                name_cell = sheet.cell(row=row, column=1)
                if name_cell.value:
                    name = str(name_cell.value).strip()
                    names.append(name)
                    
                    # 读取该行的所有数据
                    row_data = {}
                    for col, header in headers.items():
                        cell_value = sheet.cell(row=row, column=col)
                        if cell_value.value is not None:
                            row_data[header] = str(cell_value.value)
                    
                    data_rows.append(row_data)
                    total_rows += 1
            
            if not names:
                raise ValueError(f"Excel文件 '{excel_file}' 中未找到有效的数据行")
            
            print(f"📝 从 '{excel_file}' 读取到 {total_rows} 行数据，包含 {len(names)} 个名称")
            return names, data_rows, headers, excel_file
            
        except Exception as e:
            raise Exception(f"读取Excel文件 '{excel_file}' 失败: {e}")
    
    def replace_text_in_runs(self, paragraph, replacement_dict):
        """在段落的runs中替换文本（保持格式）"""
        # 首先检查整个段落是否需要替换
        full_text = paragraph.text
        needs_replacement = any(old_text in full_text for old_text in replacement_dict.keys())
        
        if not needs_replacement:
            return
        
        # 如果整个段落只需要简单替换，直接替换
        if len(paragraph.runs) == 1:
            run = paragraph.runs[0]
            for old_text, new_text in replacement_dict.items():
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
            return
        
        # 对于有多个runs的复杂段落，需要更精细的处理
        paragraph_text = full_text
        for old_text, new_text in replacement_dict.items():
            if old_text in paragraph_text:
                paragraph_text = paragraph_text.replace(old_text, new_text)
        
        # 清空原有runs并添加新文本
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = paragraph_text
    
    def replace_text_in_paragraphs(self, doc, replacement_dict):
        """替换文档正文段落中的文本"""
        for paragraph in doc.paragraphs:
            self.replace_text_in_runs(paragraph, replacement_dict)
    
    def replace_text_in_tables(self, doc, replacement_dict):
        """替换文档表格中的文本"""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_text_in_runs(paragraph, replacement_dict)
    
    def replace_text_in_headers(self, doc, replacement_dict):
        """替换文档所有页眉中的文本"""
        for section in doc.sections:
            # 处理各种页眉
            headers = [
                section.header, 
                section.first_page_header, 
                section.even_page_header
            ]
            
            for header in headers:
                if header is not None:
                    for paragraph in header.paragraphs:
                        self.replace_text_in_runs(paragraph, replacement_dict)
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    self.replace_text_in_runs(paragraph, replacement_dict)
    
    def replace_text_in_footers(self, doc, replacement_dict):
        """替换文档所有页脚中的文本"""
        for section in doc.sections:
            # 处理各种页脚
            footers = [
                section.footer, 
                section.first_page_footer, 
                section.even_page_footer
            ]
            
            for footer in footers:
                if footer is not None:
                    for paragraph in footer.paragraphs:
                        self.replace_text_in_runs(paragraph, replacement_dict)
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    self.replace_text_in_runs(paragraph, replacement_dict)
    
    def replace_text_in_document(self, doc, replacement_dict, doc_name):
        """替换Word文档中的所有文本内容"""
        print(f"🔧 开始处理文档 '{doc_name}'...")
        
        # 替换主要内容
        self.replace_text_in_paragraphs(doc, replacement_dict)
        self.replace_text_in_tables(doc, replacement_dict)
        
        # 替换页眉页脚
        self.replace_text_in_headers(doc, replacement_dict)
        self.replace_text_in_footers(doc, replacement_dict)
        
        print(f"✅ 文档 '{doc_name}' 内容替换完成")
    
    def generate_replacement_dict(self, headers, row_data):
        """生成替换字典，将表头转换为 {字段名} 格式"""
        replacement_dict = {}
        for header, value in row_data.items():
            # 创建多种格式的占位符以提高兼容性
            placeholder_variants = [
                "{" + header + "}",
                "{{" + header + "}}",
                "[" + header + "]",
                "<" + header + ">"
            ]
            
            for placeholder in placeholder_variants:
                replacement_dict[placeholder] = value
        
        return replacement_dict
    
    def process_single_excel_file(self, excel_file, docx_files):
        """处理单个Excel文件"""
        try:
            # 读取Excel数据
            names, data_rows, headers, excel_filename = self.read_excel_data(excel_file)
            
            excel_success_count = 0
            processed_files = []
            
            # 为每个数据行处理所有文档
            for data_index, (name, row_data) in enumerate(zip(names, data_rows), 1):
                try:
                    # 生成替换字典
                    replacement_dict = self.generate_replacement_dict(headers, row_data)
                    
                    print(f"🔄 处理第 {data_index} 行数据: {name}")
                    print(f"  替换映射: {list(replacement_dict.keys())}")
                    
                    # 处理所有Word文档
                    for doc_file in docx_files:
                        try:
                            # 生成新文件名
                            doc_name = Path(doc_file).stem
                            new_filename = f"{doc_name}_{name}.docx"
                            new_filepath = os.path.join(self.output_dir, new_filename)
                            
                            # 复制文档
                            shutil.copy2(doc_file, new_filepath)
                            
                            # 打开新文档进行内容替换
                            doc = Document(new_filepath)
                            
                            # 执行文本替换
                            self.replace_text_in_document(doc, replacement_dict, doc_name)
                            
                            # 保存文档
                            doc.save(new_filepath)
                            
                            processed_files.append(new_filename)
                            excel_success_count += 1
                            
                            print(f"   ✅ 生成: {new_filename}")
                            
                        except Exception as e:
                            print(f"   ❌ 处理文档 '{doc_file}' 时出错: {e}")
                            if os.path.exists(new_filepath):
                                os.remove(new_filepath)
                            continue
                    
                    print("   " + "-" * 40)
                    
                except Exception as e:
                    print(f"❌ 处理第 {data_index} 行数据时出错: {e}")
                    continue
            
            return excel_success_count, processed_files
            
        except Exception as e:
            print(f"❌ 处理Excel文件 '{excel_file}' 时发生错误: {e}")
            return 0, []
    
    def process_documents(self):
        """主处理函数"""
        try:
            # 验证文件
            docx_files, excel_files = self.validate_files()
            
            # 创建输出目录
            self.create_output_directory()
            
            total_success_count = 0
            all_processed_files = []
            
            # 处理每个Excel文件
            for excel_file in excel_files:
                print(f"\n{'='*60}")
                print(f"📋 开始处理Excel文件: {excel_file}")
                print(f"{'='*60}")
                
                success_count, processed_files = self.process_single_excel_file(excel_file, docx_files)
                total_success_count += success_count
                all_processed_files.extend(processed_files)
                
                print(f"📊 Excel文件 '{excel_file}' 处理完成，成功生成 {success_count} 个文档")
            
            # 输出处理结果
            print(f"\n{'='*60}")
            print(f"🎉 所有处理完成！")
            print(f"{'='*60}")
            print(f"总成功生成: {total_success_count} 个文档")
            
            if all_processed_files:
                print(f"\n📄 生成的文件列表:")
                for i, file in enumerate(all_processed_files, 1):
                    print(f"  {i:2d}. {file}")
            
            return total_success_count
            
        except Exception as e:
            print(f"❌ 处理过程中发生错误: {e}")
            return 0

def main():
    """主函数"""
    print("=== 智能文档批量处理工具 ===")
    print("功能: 自动处理所有.docx和.xlsx文件")
    print("=" * 60)
    print("📂 自动检测:")
    print("  - 当前目录下的所有.docx文档")
    print("  - 当前目录下的所有.xlsx文件")
    print("  - Excel中的所有数据行")
    print("=" * 60)
    
    # 创建处理器实例
    processor = DocumentProcessor()
    
    # 执行处理
    result = processor.process_documents()
    
    if result > 0:
        print(f"\n✨ 所有处理已完成！请查看 '{processor.output_dir}' 目录")
        print("\n💡 处理规则:")
        print("  - 每个Excel文件的每一行数据")
        print("  - 都会应用到每个Word文档")
        print("  - 生成: 原文档名_Excel第一列内容.docx")
    else:
        print(f"\n💥 处理失败，请检查文件格式和内容")

if __name__ == "__main__":
    main()
